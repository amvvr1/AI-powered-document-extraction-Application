import PyPDF2
import pytesseract
from PIL import Image
import docx
from pathlib import Path
from typing import Optional
import io



class DocumentReader:
    def __init__(self):
        pass

    def read_document(self, file_path: str) -> str:
        """Main entry point for document reading"""
        try:
            path = Path(file_path)
            extension = path.suffix.lower()
            
            if extension == '.pdf':
                return self._read_pdf(file_path)
            elif extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._read_image(file_path)
            elif extension == '.txt':
                return self._read_text(file_path)
            elif extension in ['.doc', '.docx']:
                return self._read_word(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
                
        except Exception as e:
            return f"Error reading document: {str(e)}"
        

    def _read_pdf(self, file_path: str) -> str:
        """Extract text from PDF files with OCR fallback"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():  # If text found
                        text += page_text + "\n"
                    else:  # No text, use OCR
                        ocr_text = self._ocr_pdf_page(file_path, page_num)
                        text += ocr_text + "\n"
        except Exception as e:
            return f"PDF reading error: {str(e)}"
        
        return text.strip() if text.strip() else "No text found in PDF"
    

    def _ocr_pdf_page(self, file_path: str, page_num: int) -> str:
        """Extract text from PDF page using OCR"""
        try:
            import fitz  # PyMuPDF
            
            # Open PDF and convert page to image
            doc = fitz.open(file_path)
            page = doc[page_num]
            pix = page.get_pixmap()
            img_data = pix.tobytes("ppm")
            
            # Convert to PIL Image and OCR
            image = Image.open(io.BytesIO(img_data))
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(image, config=custom_config)
            
            doc.close()
            return text.strip()
            
        except ImportError:
            return "[OCR failed: PyMuPDF not installed]"
        except Exception as e:
            return f"[OCR error on page {page_num + 1}: {str(e)}]"
        
    def _read_image(self, file_path: str) -> str:
        """Extract text from images using OCR"""
        try:
            image = Image.open(file_path)
            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # OCR with configuration for better accuracy
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(image, config=custom_config)
            
            return text.strip() if text.strip() else "No text detected in image"
            
        except Exception as e:
            return f"Image OCR error: {str(e)}"
        
    def _read_text(self, file_path: str) -> str:
        """Read plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                return f"Text file reading error: {str(e)}"
        except Exception as e:
            return f"Text file reading error: {str(e)}"
    
    def _read_word(self, file_path: str) -> str:
        """Extract text from Word documents"""
        try:
            doc = docx.Document(file_path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text) if text else "No text found in Word document"
            
        except Exception as e:
            return f"Word document reading error: {str(e)}"


    def get_supported_formats(self) -> list:
        """Return list of supported file formats"""
        return ['.pdf', '.txt', '.doc', '.docx', '.png', '.jpg', '.jpeg', '.tiff', '.bmp']
    
    def validate_file(self, file_path: str) -> bool:
        """Check if file format is supported"""
        extension = Path(file_path).suffix.lower()
        return extension in self.get_supported_formats()


if __name__ == "__main__":
    # Test the DocumentReader class
    reader = DocumentReader()
    
    print("DocumentReader initialized successfully!")
    print(f"Supported formats: {reader.get_supported_formats()}")
    
   
   
